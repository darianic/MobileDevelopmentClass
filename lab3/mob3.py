from docx import Document
from docx.shared import Pt
from docx2pdf import convert


wdFormatPDF = 17

doc = Document('correct.docx')

#Результаты лаб.1/лаб.2
with open("internet.txt") as file:
    internet = file.read()


with open("sms.txt") as file:
    sms = file.read()


with open("telephony.txt") as file:
    telephony = file.read()


table1 = doc.tables[0]
table2 = doc.tables[1]
table3 = doc.tables[2]
table4 = doc.tables[3]
table5 = doc.tables[4]
table6 = doc.tables[5]
table7 = doc.tables[6]

#Первая таблица
bank = table1.rows[0].cells[0]
receiver = table1.rows[3].cells[0]
inn = table1.rows[2].cells[1]
kpp = table1.rows[2].cells[3]
bik = table1.rows[0].cells[5]
sch1 = table1.rows[1].cells[5]
sch2 = table1.rows[2].cells[5]

bank.text = 'АО "Банк Лабораторная" г. Санкт-Петербург'
receiver.text = 'ООО "Сдана"'
inn.text = '7727737666'
kpp.text = '777777001'
bik.text = '0445255703'
sch1.text = '30101810200000000300'
sch2.text = '40337770000000003822'

#Вторая таблица
n_sch = table2.rows[0].cells[2]
data = table2.rows[0].cells[4]

n_sch.text = '84'
data.text = '30.04.2020'

for row in table2.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        paragraph = paragraphs[0]
        run_obj = paragraph.runs
        run = run_obj[0]
        font = run.font
        font.size = Pt(14)
        font.bold = True

#Третья таблица
provider = table3.rows[0].cells[1]
buyer = table3.rows[1].cells[1]
base = table3.rows[2].cells[1]

provider.text = 'ООО "Сдана"'+', '+'ИНН: 7727737666'+', '+'КПП: 777777001'+', '+'191002, г.Санкт-Петербург, улица Ломоносова, 9'
buyer.text = 'ООО "Дарья"'+', '+'ИНН: 1234567890'+', '+'КПП: 0987654321'+', ''197101, г.Санкт-Петербург, Кронверкский проспект, 49'
base.text = '№97757910982' + ' от' + ' 30.04.2020'

#Четвертая таблица
num = table4.rows[1].cells[0]
work_name = table4.rows[1].cells[1]
count = table4.rows[1].cells[2]
ed = table4.rows[1].cells[3]
price = table4.rows[1].cells[4]
sum = table4.rows[1].cells[5]

#Для количества наименований
stroka = 1

num.text = '1'
work_name.text = 'Плата за услуги Интернет'
count.text = '1'
ed.text = 'шт'
price.text = '0.5'
sum.text = str(internet)


doc.tables[3].add_row()
stroka = stroka + 1

num = table4.rows[2].cells[0]
work_name = table4.rows[2].cells[1]
count = table4.rows[2].cells[2]
ed = table4.rows[2].cells[3]
price = table4.rows[2].cells[4]
sum = table4.rows[2].cells[5]

num.text = '2'
work_name.text = 'Плата за услуги Телефония'
count.text = '1'
ed.text = 'шт'
price.text = '1'
sum.text = str(telephony)

s2 = float(sum.text)

doc.tables[3].add_row()
stroka = stroka + 1

num = table4.rows[3].cells[0]
work_name = table4.rows[3].cells[1]
count = table4.rows[3].cells[2]
ed = table4.rows[3].cells[3]
price = table4.rows[3].cells[4]
sum = table4.rows[3].cells[5]

num.text = '3'
work_name.text = 'Плата за услуги SMS'
count.text = '1'
ed.text = 'шт'
price.text = '1'
sum.text = str(sms)

s3 = float(sum.text)

#Пятая таблица
itog = table5.rows[0].cells[1]
nds = table5.rows[1].cells[1]
vsego = table5.rows[2].cells[1]

s = round(float(internet) + float(sms) + float(telephony), 2)

itog.text = str(s)
nds.text = str(round(s * 18 / 100, 2))
vsego.text = str(s)

#Шестая таблица
all_names = table6.rows[0].cells[1]

all_names.text = str(stroka) + ', ' + 'на сумму ' + str(s)


supervisor = table7.rows[0].cells[1]
accountant = table7.rows[0].cells[3]

supervisor.text = 'Ниценкова Д. В.'
accountant.text = 'Ниценкова Д. В.'


doc.save('result.docx')

print('Converting docx to pdf:')
convert('result.docx', 'res.pdf')